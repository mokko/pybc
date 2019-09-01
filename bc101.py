from docx import Document
from docx.shared import Cm
import re
import sys
from glob import iglob
import os
import time
from PIL import Image

VERBOSE=1
MONITORDIR='C:/Users/User/eclipse-workspace/py-bc' #DOnt end with slash

def error (msg):
    print ('ERROR: '+ msg)
    sys.exit(1)

def verbose (msg):
    if VERBOSE:
        print (msg)

def write_bc (text,format):
    if text != '':
        verbose ('TT%s'% text)
        if format == '1':
            import barcode
            from barcode import generate
            from barcode.writer import ImageWriter
            code=barcode.get('code128', text, writer=ImageWriter())

            #code.save ('temp', options={
            image=code.render (writer_options={
                'module_width':0.2, #default 0.2 in mm 
                'module_height':3,   #default 15 in mm
                'quiet_zone':1,    #default 6.5
                'font-size':3,     #default 10 (integer)
                'text_distance':3.0,  #default 5
                'write_text': False, #default True
            })
            #image=image.resize((int(image.size[0]*0.6),40)) #resample=Image.LANCZOS
            print (image.size)
            size=[int(x*0.7) for x in image.size] 
            print (size)
            #image.thumbnail((250,60), Image.LANCZOS) # resizes image in place
            image.thumbnail(size)
            image.save('temp.png', dpi=(100,100))
            
        elif format == '2':
            #requires ghostscript
            import treepoem
            image = treepoem.generate_barcode(barcode_type='code128',data=text)
            #image=image.convert('1')
            #image=image.resize((136,20), resample=Image.LANCZOS)
            #image.thumbnail((100,40), Image.LANCZOS) # resizes image in place
            #image=image.rotate(90, expand=True,  fillcolor="white")
            image.save('temp.png')
            #print (image)
        elif format == '3':
            from pubcode import Code128
            #print (len(text))
            barcode=Code128(text, charset='A')
            image=barcode.image()# use defaults and do resize on our own
            #print (image.size[0]) 
            image=image.resize((image.size[0],20)) #resample=Image.LANCZOS
            image.save('temp.png')
        else:
            error ("No format recognized!")


def transform (infile, outfile):    
    doc = Document(infile)  # input
    ncols=len(doc.tables[0].columns) # only work on first table
    nrows=len(doc.tables[0].rows)
    verbose ("Table 0 grid: %i/%i"  % (nrows, ncols))
    
    marked_cells={}
    
    for rid in range (0,nrows):
        for cid in range (0, ncols):
            cell=doc.tables[0].rows[rid].cells[cid].text
            if rid == 0:
                m=re.match('{(\d)}', cell)
                if m:
                    #leave it as is, as it messes up the formatting
                    #doc.tables[0].rows[rid].cells[cid].text=cell.split('}')[1]
                    marked_cells[cid]=m.group(1)
            else:
                if cid in marked_cells:
                    write_bc (cell, marked_cells[cid])
                    doc.tables[0].rows[rid].cells[cid].text=''
                    p = doc.tables[0].rows[rid].cells[cid].add_paragraph()
                    r = p.add_run()
                    r.add_picture('temp.png')                    
                    
    #print (marked_cells)
    doc.save(outfile)


def print_cursor ():
    global c
    if c==0:
        symbol='.'
        c=1
    else:
        symbol='+'
        c=0            
    print('\r'+symbol, end='', flush=True)


if __name__ == '__main__': 
    verbose ('Listening for changes in %s' % MONITORDIR)
    c=0
    while True:
        for fn in iglob(MONITORDIR+'/*.docx'):
            if os.path.isfile(fn): # not dir, can be file or link
                base=os.path.basename(fn[:-5])
                m=re.search('-bc$', base)
                if not m:
                    outfile=base+'-bc.docx'
                    if not os.path.exists(outfile): 
                        verbose ('%s --> %s' % (fn, outfile))
                        transform(fn, outfile)
                    #else:
                    #    verbose ('Not overwritting %s' %outfile)
        time.sleep(3)    
        print_cursor()

#
#

