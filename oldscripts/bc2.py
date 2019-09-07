import barcode
from barcode import generate
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Cm
import re

doc = Document('Maurice.docx')

ncols=len(doc.tables[0].columns)
nrows=len(doc.tables[0].rows)
print ("ncols %s" % ncols)
print ("nrows %s" % nrows)

marked_cells={}

for rid in range (0,nrows):
    for cid in range (0, ncols):
        cell=doc.tables[0].rows[rid].cells[cid].text
        if rid == 0:
            m=re.match('{(\d)}', cell)
            if m:
                #m = re.match("(\w+\ )",str)
                marked_cells[cid]=m.group(1)
        else:
            if cid in marked_cells:
                print ('H%s'%cell)
print (marked_cells)
doc.save('newdoc.docx')


#
#

