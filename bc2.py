import barcode
from barcode import generate
from barcode.writer import ImageWriter
from docx import Document
from docx.shared import Cm

doc = Document('Objektliste mit Barcode.docx')


for row in doc.tables[0].rows:
    if (row.cells[2].text != doc.tables[0].rows[0].cells[2].text):
        objId=int(row.cells[2].text) #test if not integer?
        row.cells[2].text=''
        format="OB%09d" % objId
        code=barcode.get('code128', format, writer=ImageWriter())
        code.save(format+'-A')
        code.save(format+'-B', options={'module_width':0.1, 'module_height':5})
        code.save(format+'-C', options={
            'module_width':.4, 
            'module_height':15.0, 
            'quiet_zone':6.5, 
            'font_size':10, 
            'text_distance':5.0,
            'text':''
            'background':'white',
            'foreground':'black',
        })
        p = row.cells[2].add_paragraph()
        r = p.add_run()
        r.add_picture(format+'-A.png')
doc.save('newdoc.docx')


#
#

